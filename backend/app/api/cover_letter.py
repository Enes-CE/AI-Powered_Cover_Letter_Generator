from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import FileResponse
from app.models.schemas import (
    CoverLetterRequest,
    CoverLetterResponse,
    JobAnalysis,
    SkillMatch,
    ToneType,
    CoverLetterBatchResponse,
    ExportRequest,
)
from app.services.spacy_service import SpaCyService
from app.services.ollama_service import OllamaAiService
from app.settings import settings
from typing import List, Union
import io
import tempfile
import os
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

router = APIRouter()
nlp_service = SpaCyService()
ai_service = OllamaAiService()

@router.post("/generate-cover-letter", response_model=Union[CoverLetterResponse, CoverLetterBatchResponse])
async def generate_cover_letter(request: CoverLetterRequest):
    """
    Generate a personalized cover letter based on job posting and CV
    """
    try:
        # Extract job information
        job_info = nlp_service.extract_job_info(request.job_posting.job_posting_text)
        
        # Extract skills from job posting
        job_skills = nlp_service.extract_skills_from_text(request.job_posting.job_posting_text)
        key_requirements = nlp_service.extract_key_requirements(request.job_posting.job_posting_text)
        
        # Extract skills from CV
        cv_skills = nlp_service.analyze_cv_skills(request.cv_data.cv_text or "")
        
        # Match skills
        skill_matches = nlp_service.match_skills(job_skills, cv_skills)
        
        # Find missing skills
        missing_skills = nlp_service.find_missing_skills(job_skills, cv_skills)
        
        # Generate recommendations
        recommendations = nlp_service.generate_recommendations(skill_matches, missing_skills)
        
        # Create analysis
        analysis = JobAnalysis(
            extracted_skills=job_skills,
            required_experience=job_info.get('required_experience', '3+ years'),
            company_name=job_info.get('company_name', 'Tech Company'),
            position_title=job_info.get('position_title', 'Software Engineer'),
            key_requirements=key_requirements
        )
        
        # Single or multi-variant generation
        num_variants = max(1, int(request.variants or 1))
        tones_cycle = ['formal', 'friendly', 'concise']
        letters: List[str] = []
        tones_used: List[str] = []
        for i in range(num_variants):
            tone_to_use = request.tone
            if num_variants > 1:
                tone_to_use = tones_cycle[i % len(tones_cycle)]
            if settings.AI_PROVIDER == "ollama":
                letter = await ai_service.draft_cover_letter(
                    job_info=job_info,
                    cv_skills=cv_skills,
                    skill_matches=skill_matches,
                    tone=tone_to_use,
                )
            else:
                letter = generate_cover_letter(job_info, cv_skills, skill_matches, tone_to_use)
            letters.append(letter)
            tones_used.append(tone_to_use)
        
        # Convert skill matches to SkillMatch objects
        skill_match_objects = [
            SkillMatch(
                skill=match['skill'],
                matched=match['matched'],
                confidence=match['confidence'],
                cv_evidence=match['cv_evidence']
            )
            for match in skill_matches
        ]
        
        if num_variants == 1:
            return CoverLetterResponse(
                cover_letter=letters[0],
                analysis=analysis,
                skill_matches=skill_match_objects,
                missing_skills=missing_skills,
                recommendations=recommendations,
                tone_used=tones_used[0],
            )
        else:
            return CoverLetterBatchResponse(
                letters=letters,
                analysis=analysis,
                skill_matches=skill_match_objects,
                missing_skills=missing_skills,
                recommendations=recommendations,
                tone_used=tones_used,
            )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating cover letter: {str(e)}")

@router.post("/export-pdf")
async def export_cover_letter_pdf(request: ExportRequest):
    """Export cover letter as PDF"""
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            # Create PDF document
            doc = SimpleDocTemplate(tmp_file.name, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leading=14
            )
            
            # Build PDF content
            story = []
            
            # Title
            story.append(Paragraph(f"Cover Letter - {request.job_title}", title_style))
            story.append(Spacer(1, 20))
            
            # Company info
            story.append(Paragraph(f"<b>Company:</b> {request.company_name}", normal_style))
            story.append(Paragraph(f"<b>Position:</b> {request.job_title}", normal_style))
            story.append(Spacer(1, 20))
            
            # Cover letter content
            paragraphs = request.cover_letter.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), normal_style))
                    story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            
            # Return the file
            return FileResponse(
                tmp_file.name,
                media_type='application/pdf',
                filename=f'cover_letter_{request.company_name.replace(" ", "_")}_{request.job_title.replace(" ", "_")}.pdf'
            )
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")

@router.post("/export-docx")
async def export_cover_letter_docx(request: ExportRequest):
    """Export cover letter as DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading(f'Cover Letter - {request.job_title}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Company info
            doc.add_paragraph(f'Company: {request.company_name}')
            doc.add_paragraph(f'Position: {request.job_title}')
            doc.add_paragraph('')  # Empty line
            
            # Cover letter content
            paragraphs = request.cover_letter.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Save document
            doc.save(tmp_file.name)
            
            # Return the file
            return FileResponse(
                tmp_file.name,
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                filename=f'cover_letter_{request.company_name.replace(" ", "_")}_{request.job_title.replace(" ", "_")}.docx'
            )
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")

@router.get("/test")
async def test_endpoint():
    """Test endpoint to verify API is working"""
    return {"message": "Cover letter API is working!", "status": "success"}


def generate_cover_letter(job_info: dict, cv_skills: List[str], skill_matches: List[dict], tone: str) -> str:
    """Generate cover letter based on job info, CV skills, and tone"""
    
    position_title = job_info.get('position_title', 'Software Engineer')
    company_name = job_info.get('company_name', 'Tech Company')
    
    # Get matched skills
    matched_skills = [match['skill'] for match in skill_matches if match['matched']]
    top_skills = matched_skills[:3] if matched_skills else ['software development']
    
    # Generate based on tone
    if tone == "formal":
        cover_letter = f"""
Dear Hiring Manager,

I am writing to express my strong interest in the {position_title} position at {company_name}. With my background in software development and experience with {', '.join(top_skills)}, I believe I would be a valuable addition to your team.

My experience includes:
- {len(cv_skills)} technical skills including {', '.join(top_skills)}
- Strong problem-solving and analytical abilities
- Experience with modern development practices
- Collaborative team environment experience

I am excited about the opportunity to contribute to your innovative projects and grow with your team.

Best regards,
[Your Name]
        """.strip()
    
    elif tone == "friendly":
        cover_letter = f"""
Hi there!

I'm really excited about the {position_title} opportunity at {company_name}! I think my background in {', '.join(top_skills)} would be a great fit for your team.

Here's what I bring to the table:
- Solid experience with {', '.join(top_skills)}
- A passion for learning new technologies
- Great teamwork and communication skills
- A track record of delivering quality results

I'd love to chat about how I can contribute to your team's success!

Best,
[Your Name]
        """.strip()
    
    else:  # concise
        cover_letter = f"""
Dear Hiring Manager,

I'm interested in the {position_title} position at {company_name}.

Key qualifications:
- {len(cv_skills)} technical skills
- Experience with {', '.join(top_skills)}
- Strong development background

Available for immediate start.

Best regards,
[Your Name]
        """.strip()
    
    return cover_letter


@router.post("/extract-cv-text")
async def extract_cv_text(file: UploadFile = File(...)):
    """Extract plain text from an uploaded PDF file."""
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    try:
        from PyPDF2 import PdfReader
        import io

        content = await file.read()
        
        # Check if file is empty
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="PDF file is empty")
        
        # Try to read the PDF
        try:
            reader = PdfReader(io.BytesIO(content))
        except Exception as pdf_error:
            raise HTTPException(status_code=400, detail=f"Invalid PDF format: {str(pdf_error)}")
        
        # Check if PDF has pages
        if len(reader.pages) == 0:
            raise HTTPException(status_code=400, detail="PDF has no readable pages")
        
        pages_text = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                if text.strip():  # Only add non-empty text
                    pages_text.append(text)
            except Exception as page_error:
                # Skip problematic pages but continue with others
                continue
        
        if not pages_text:
            raise HTTPException(status_code=400, detail="No text could be extracted from PDF. The PDF might be image-based or corrupted.")
        
        extracted_text = "\n".join(pages_text).strip()
        
        # Check if we got meaningful text
        if len(extracted_text) < 10:
            raise HTTPException(status_code=400, detail="Extracted text is too short. The PDF might be image-based.")
        
        return {"text": extracted_text}
        
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse PDF: {str(e)}")
